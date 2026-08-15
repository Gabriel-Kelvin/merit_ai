import io
import time

import pytest
from docx import Document

from app.providers import call_with_deadline
from app.resumes.models import ResumeProfile
from app.resumes.parser import ResumeFormatError, _recover_structured_context, extract_resume_text


def test_provider_call_has_a_true_wall_clock_deadline():
    started = time.monotonic()

    with pytest.raises(TimeoutError):
        call_with_deadline(lambda: time.sleep(0.2), 0.02)

    assert time.monotonic() - started < 0.1


def test_extracts_readable_txt_resume():
    text = extract_resume_text(
        "candidate.txt",
        b"Maya Singh\nB.Tech Computer Science\nBuilt an API using Python and FastAPI.",
    )

    assert "Maya Singh" in text
    assert "Python and FastAPI" in text


def test_extracts_docx_paragraphs_and_table_cells():
    document = Document()
    document.add_paragraph("Maya Singh - Software Engineer")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Skills"
    table.cell(0, 1).text = "Python, FastAPI, React"
    buffer = io.BytesIO()
    document.save(buffer)

    text = extract_resume_text("candidate.docx", buffer.getvalue())

    assert "Software Engineer" in text
    assert "Python, FastAPI, React" in text


def test_rejects_unsupported_and_unreadable_resumes():
    with pytest.raises(ResumeFormatError, match="PDF, DOCX, and TXT"):
        extract_resume_text("candidate.rtf", b"This content is long enough but unsupported.")

    with pytest.raises(ResumeFormatError, match="too little readable text"):
        extract_resume_text("candidate.txt", b"short")


def test_recovers_work_and_certification_context_when_ai_omits_them():
    text = """Maya Singh
WORK EXPERIENCE
Software Engineer | Acme Labs | 2023 - Present
Built a FastAPI service and reduced routing time by 40%.
EDUCATION
B.Tech Computer Science, 2022
CERTIFICATIONS
AWS Certified Developer
"""
    profile = _recover_structured_context(
        ResumeProfile(name="Maya Singh", technical_skills=["Python", "FastAPI"]),
        text,
    )

    assert profile.work_experience[0].title == "Software Engineer"
    assert profile.work_experience[0].company == "Acme Labs"
    assert "reduced routing time by 40%" in profile.work_experience[0].achievements[0]
    assert profile.certifications == ["AWS Certified Developer"]
